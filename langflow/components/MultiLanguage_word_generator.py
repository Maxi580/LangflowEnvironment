import json
from langflow.custom import Component
from langflow.io import Output, HandleInput, MultilineInput
from langflow.schema.message import Message
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import base64
import os
from dotenv import load_dotenv

load_dotenv()

DOCX_MAGIC_BYTES = os.getenv("DOCX_MAGIC_BYTES")


class MultiLanguageWordComponent(Component):
    display_name = "Multi-Language Word Document"
    description = (
        "Takes the Translation Engine JSON and creates one Word document per language. "
        "Each DOCX contains all translated sections in a structured format."
    )
    icon = "📄"

    inputs = [
        HandleInput(
            name="translations_json",
            display_name="Translations JSON",
            info="The JSON output from the Translation Engine.",
            input_types=["Message"],
            is_list=False,
        ),
    ]

    outputs = [
        Output(
            display_name="Response",
            name="response_output",
            method="create_word_documents",
        ),
    ]

    # ------------------------------------------------------------------ #
    #  Section display config                                              #
    # ------------------------------------------------------------------ #

    # Maps Translation Engine keys to readable section titles in the DOCX
    SECTION_ORDER = [
        ("About_Client", "About Client"),
        ("Challenge", "Challenge"),
        ("Solution", "Solution"),
        ("Why_Atos", "Why Us"),
        ("Business_Impact", "Business Impact"),
        ("IT_Impact", "IT Impact"),
        ("Technology", "Technology"),
    ]

    # ------------------------------------------------------------------ #
    #  Helpers                                                             #
    # ------------------------------------------------------------------ #

    def _add_section(self, doc, title: str, content: str):
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(14)

        content_paragraph = doc.add_paragraph()
        content_run = content_paragraph.add_run(content)
        content_run.font.size = Pt(11)

        doc.add_paragraph()

    def _create_single_docx(self, language: str, sections: dict) -> tuple[bytes, str] | None:
        """Create one DOCX for the given language. Returns (file_bytes, filename) or None."""
        try:
            doc = Document()

            # Title
            customer = sections.get("Client_Name", "Client")
            project = sections.get("Project", "Project")

            title_p = doc.add_paragraph()
            title_run = title_p.add_run(f"Reference Case: {customer}")
            title_run.bold = True
            title_run.font.size = Pt(16)
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            project_p = doc.add_paragraph()
            project_run = project_p.add_run(f"({project})")
            project_run.bold = True
            project_run.font.size = Pt(14)
            project_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Language indicator
            lang_p = doc.add_paragraph()
            lang_run = lang_p.add_run(f"Language: {language}")
            lang_run.font.size = Pt(10)
            lang_run.italic = True
            lang_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()
            doc.add_paragraph()

            # Sections
            for section_key, section_title in self.SECTION_ORDER:
                content = sections.get(section_key, "")
                if content:
                    self._add_section(doc, section_title, content)

            # Save to bytes
            tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
            tmp_path = tmp.name
            tmp.close()

            doc.save(tmp_path)
            with open(tmp_path, "rb") as f:
                file_bytes = f.read()
            os.unlink(tmp_path)

            # Filename
            customer_slug = customer.replace(" ", "_").lower()
            lang_slug = language.replace(" ", "_").lower()
            filename = f"reference_{customer_slug}_{lang_slug}.docx"

            return file_bytes, filename

        except Exception as e:
            print(f"Error creating DOCX for {language}: {e}")
            return None

    # ------------------------------------------------------------------ #
    #  Main output                                                         #
    # ------------------------------------------------------------------ #

    def create_word_documents(self) -> Message:
        # Parse translations JSON
        raw = ""
        if isinstance(self.translations_json, Message):
            raw = self.translations_json.text or ""
        elif isinstance(self.translations_json, str):
            raw = self.translations_json
        else:
            raw = str(self.translations_json)

        if not raw.strip():
            return Message(text="No translation data received.")

        try:
            data = json.loads(raw)
        except json.JSONDecodeError as e:
            return Message(text=f"Failed to parse translations JSON: {e}")

        translations = data.get("translations", [])
        if not translations:
            return Message(text="No translations found in JSON.")

        # Generate one DOCX per language
        response_parts = []
        files_created = 0

        for entry in translations:
            language = entry.get("language", "Unknown")
            sections = entry.get("sections", {})
            error = entry.get("error")

            if error:
                response_parts.append(f"⚠️ Skipped {language}: {error}")
                continue

            if not sections:
                response_parts.append(f"⚠️ Skipped {language}: no sections.")
                continue

            result = self._create_single_docx(language, sections)
            if result is None:
                response_parts.append(f"❌ Failed to create DOCX for {language}.")
                continue

            file_bytes, filename = result
            b64_content = base64.b64encode(file_bytes).decode("utf-8")

            response_parts.append(
                f"{language}: {filename}\n\n"
                f"<{DOCX_MAGIC_BYTES}>\n"
                f"filename:{filename}\n"
                f"content_type:application/vnd.openxmlformats-officedocument.wordprocessingml.document\n"
                f"size:{len(file_bytes)}\n"
                f"data:{b64_content}\n"
                f"</{DOCX_MAGIC_BYTES}>"
            )
            files_created += 1

        self.status = f"Created {files_created} DOCX file(s) for {len(translations)} language(s)."

        return Message(text="\n\n".join(response_parts))