from langflow.custom import Component
from langflow.io import Output, MultilineInput
from langflow.schema import Data
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import base64
import os
from dotenv import load_dotenv

load_dotenv()

DOCX_MAGIC_BYTES = os.getenv("DOCX_MAGIC_BYTES")


class AtosTemplateWordComponent(Component):
    display_name = "Atos Template Word Document"
    description = "Creates Word document with structured sections for client information"
    icon = "📄"
    inputs = [
        MultilineInput(
            name="customer_name",
            display_name="Customer Name",
            info="Customer name for the document filename",
            required=True,
        ),
        MultilineInput(
            name="about_client",
            display_name="About Client",
            info="Content for the About Client section",
            required=True,
        ),
        MultilineInput(
            name="challenge_text",
            display_name="Challenge Description",
            info="Content for the Challenge section",
            required=True,
        ),
        MultilineInput(
            name="solution_text",
            display_name="Solution Description",
            info="Content for the Solution section",
            required=True,
        ),
        MultilineInput(
            name="why_us_text",
            display_name="Why Us",
            info="Content for the Why Us section",
            required=True,
        ),
        MultilineInput(
            name="business_impact_text",
            display_name="Business Impact",
            info="Content for the Business Impact section",
            required=True,
        ),
        MultilineInput(
            name="it_impact_text",
            display_name="IT Impact",
            info="Content for the IT Impact section",
            required=True,
        ),
        MultilineInput(
            name="technology_text",
            display_name="Technology",
            info="Content for the Technology section",
            required=True,
        ),
    ]

    outputs = [
        Output(
            display_name="Response",
            name="response_output",
            method="create_word_document"
        )
    ]

    def add_section(self, doc, title, content):
        """
        Add a section with bold title and content to the document
        """
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(14)

        content_paragraph = doc.add_paragraph()
        content_run = content_paragraph.add_run(content)
        content_run.font.size = Pt(11)

        doc.add_paragraph()


    def create_word_document(self) -> Data:
        """Create Word document with structured sections"""

        try:
            doc = Document()

            title_paragraph = doc.add_paragraph()
            title_run = title_paragraph.add_run(f"Reference Case: {self.customer_name}")
            title_run.bold = True
            title_run.font.size = Pt(16)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add some space after title
            doc.add_paragraph()
            doc.add_paragraph()

            # Define sections with their content
            sections = [
                ("About Client", self.about_client),
                ("Challenge", self.challenge_text),
                ("Solution", self.solution_text),
                ("Why Us", self.why_us_text),
                ("Business Impact", self.business_impact_text),
                ("IT Impact", self.it_impact_text),
                ("Technology", self.technology_text),
            ]

            for title, content in sections:
                self.add_section(doc, title, content)

            print(f"✅ Document created with {len(sections)} sections")

            temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
            temp_file_path = temp_file.name
            temp_file.close()

            doc.save(temp_file_path)

            with open(temp_file_path, 'rb') as f:
                file_content = f.read()

            os.unlink(temp_file_path)

            base64_content = base64.b64encode(file_content).decode('utf-8')
            filename = f"reference_{self.customer_name.replace(' ', '_').lower()}.docx"

            text_response = f"""Word document created successfully!

<{DOCX_MAGIC_BYTES}>
filename:{filename}
content_type:application/vnd.openxmlformats-officedocument.wordprocessingml.document
size:{len(file_content)}
data:{base64_content}
</{DOCX_MAGIC_BYTES}>"""

            return Data(data={"text": text_response})

        except Exception as e:
            error_text = f"❌ Failed to create Word document: {str(e)}"
            return Data(data={"text": error_text})
